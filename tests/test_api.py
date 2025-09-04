from fastapi.testclient import TestClient
from app.main import app
from io import BytesIO

client = TestClient(app)

def test_analyze_endpoint_requires_file():
    r = client.post('/analyze')
    assert r.status_code == 422  # missing file

def test_analyze_with_docx(tmp_path):
    # create a small docx in memory
    from docx import Document
    p = tmp_path / 'sample.docx'
    doc = Document()
    doc.add_paragraph('This is a test.')
    doc.save(str(p))
    with open(p, 'rb') as fh:
        r = client.post('/analyze', files={'file': ('sample.docx', fh, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
    assert r.status_code == 200
    data = r.json()
    assert 'report' in data

def test_correct_endpoint_returns_zip(tmp_path):
    from docx import Document
    p = tmp_path / 'sample.docx'
    doc = Document()
    doc.add_paragraph('This is a test for correction.')
    doc.save(str(p))
    with open(p, 'rb') as fh:
        r = client.post('/correct', files={'file': ('sample.docx', fh, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
    assert r.status_code == 200
    assert r.headers['content-type'] == 'application/zip'
