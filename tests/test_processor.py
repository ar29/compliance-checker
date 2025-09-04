import pytest
from app.processor import (
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_file,
)
from docx import Document
import tempfile, os


def test_extract_docx(tmp_path):
    p = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Hello world.")
    doc.add_paragraph("This is a test.")
    doc.save(str(p))
    text = extract_text_from_docx(str(p))
    assert "Hello world" in text
    assert "This is a test" in text


def test_extract_file_unsupported(tmp_path):
    p = tmp_path / "test.txt"
    p.write_text("hi")
    with pytest.raises(ValueError):
        extract_text_from_file(str(p))
